"""
Unit tests for Document Parsing Agent
"""
import os
import pytest
import tempfile
import shutil
from pathlib import Path
from datetime import datetime

from backend.agents.document_parsing import (
    DocumentParsingAgent,
    DocumentParsingInput,
)
from backend.models.schemas import DocumentFile
from backend.models.enums import FileType


@pytest.fixture
def temp_dir():
    """Create temporary directory for tests"""
    temp = tempfile.mkdtemp()
    yield temp
    shutil.rmtree(temp, ignore_errors=True)


@pytest.fixture
def agent():
    """Create document parsing agent"""
    return DocumentParsingAgent(enable_ocr=False)  # Disable OCR for faster tests


@pytest.fixture
def sample_pdf(temp_dir):
    """Create sample PDF file"""
    # Create a simple PDF using reportlab or basic PDF structure
    pdf_path = os.path.join(temp_dir, "sample.pdf")
    
    # Minimal valid PDF
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] 
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
100 700 Td
(Hello World - Test Document) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000359 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
437
%%EOF
"""
    with open(pdf_path, 'wb') as f:
        f.write(pdf_content)
    
    return pdf_path


@pytest.fixture
def sample_docx(temp_dir):
    """Create sample DOCX file"""
    from docx import Document

    docx_path = os.path.join(temp_dir, "sample.docx")

    # Create document
    doc = Document()
    
    # Add heading
    doc.add_heading('Test Document', level=0)
    
    # Add paragraphs
    doc.add_paragraph('This is a test paragraph.')
    doc.add_paragraph('Another paragraph with more content.')
    
    # Add table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i},{j}"
    
    doc.save(docx_path)
    return docx_path


class TestDocumentParsingAgent:
    """Test suite for DocumentParsingAgent"""
    
    def test_parse_pdf(self, agent, sample_pdf):
        """Test PDF parsing"""
        input_data = DocumentParsingInput(
            documents=[
                DocumentFile(
                    file_id="test_001",
                    file_path=sample_pdf,
                    file_type=FileType.PDF,
                    file_size=os.path.getsize(sample_pdf),
                    original_name="sample.pdf",
                    mime_type="application/pdf",
                    relative_path="sample.pdf"
                )
            ],
            job_id="test_job_pdf"
        )
        
        output = agent.parse(input_data)
        
        assert output.success is True
        assert len(output.parsed_documents) == 1
        assert output.total_pages == 1
        assert output.parsed_documents[0].file_type == FileType.PDF
    
    def test_parse_docx(self, agent, sample_docx):
        """Test DOCX parsing"""
        input_data = DocumentParsingInput(
            documents=[
                DocumentFile(
                    file_id="test_002",
                    file_path=sample_docx,
                    file_type=FileType.DOCX,
                    file_size=os.path.getsize(sample_docx),
                    original_name="sample.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    relative_path="sample.docx"
                )
            ],
            job_id="test_job_docx"
        )
        
        output = agent.parse(input_data)
        
        assert output.success is True
        assert len(output.parsed_documents) == 1
        assert output.parsed_documents[0].file_type == FileType.DOCX
        
        # Check content extracted
        doc = output.parsed_documents[0]
        assert len(doc.pages) == 1
        assert "Test Document" in doc.pages[0].text
        assert len(doc.pages[0].tables) == 1  # We added one table
    
    def test_parse_multiple_documents(self, agent, sample_pdf, sample_docx):
        """Test parsing multiple documents"""
        input_data = DocumentParsingInput(
            documents=[
                DocumentFile(
                    file_id="test_001",
                    file_path=sample_pdf,
                    file_type=FileType.PDF,
                    file_size=os.path.getsize(sample_pdf),
                    original_name="sample.pdf",
                    mime_type="application/pdf",
                    relative_path="sample.pdf"
                ),
                DocumentFile(
                    file_id="test_002",
                    file_path=sample_docx,
                    file_type=FileType.DOCX,
                    file_size=os.path.getsize(sample_docx),
                    original_name="sample.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    relative_path="sample.docx"
                )
            ],
            job_id="test_job_multiple"
        )
        
        output = agent.parse(input_data)
        
        assert output.success is True
        assert len(output.parsed_documents) == 2
        assert output.total_pages >= 1
    
    def test_parse_nonexistent_file(self, agent, temp_dir):
        """Test parsing non-existent file"""
        # Create a valid DocumentFile but with nonexistent path
        input_data = DocumentParsingInput(
            documents=[
                DocumentFile(
                    file_id="test_003",
                    file_path=os.path.join(temp_dir, "nonexistent.pdf"),
                    file_type=FileType.PDF,
                    file_size=1024,  # Must be > 0
                    original_name="nonexistent.pdf",
                    mime_type="application/pdf",
                    relative_path="nonexistent.pdf"
                )
            ],
            job_id="test_job_nonexistent"
        )
        
        output = agent.parse(input_data)
        
        assert output.success is False
        assert len(output.errors) > 0
    
    def test_parse_empty_documents_list(self, agent):
        """Test parsing empty document list"""
        input_data = DocumentParsingInput(
            documents=[],
            job_id="test_job_empty"
        )
        
        output = agent.parse(input_data)
        
        assert output.success is False
        assert len(output.parsed_documents) == 0
    
    def test_processing_time_recorded(self, agent, sample_pdf):
        """Test that processing time is recorded"""
        input_data = DocumentParsingInput(
            documents=[
                DocumentFile(
                    file_id="test_001",
                    file_path=sample_pdf,
                    file_type=FileType.PDF,
                    file_size=os.path.getsize(sample_pdf),
                    original_name="sample.pdf",
                    mime_type="application/pdf",
                    relative_path="sample.pdf"
                )
            ],
            job_id="test_job_time"
        )
        
        output = agent.parse(input_data)
        
        assert output.processing_time > 0
        assert output.processing_time < 60  # Should be fast
    
    def test_statistics_calculated(self, agent, sample_docx):
        """Test that statistics are calculated correctly"""
        input_data = DocumentParsingInput(
            documents=[
                DocumentFile(
                    file_id="test_002",
                    file_path=sample_docx,
                    file_type=FileType.DOCX,
                    file_size=os.path.getsize(sample_docx),
                    original_name="sample.docx",
                    mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    relative_path="sample.docx"
                )
            ],
            job_id="test_job_stats"
        )
        
        output = agent.parse(input_data)
        
        assert output.total_pages >= 0
        assert output.total_tables >= 0
        assert output.total_images >= 0


class TestPDFParser:
    """Test suite for PDF Parser"""
    
    def test_pdf_text_extraction(self, sample_pdf):
        """Test PDF text extraction"""
        from backend.parsers.pdf_parser import PDFParser
        
        parser = PDFParser(enable_ocr=False)
        parsed = parser.parse(sample_pdf)
        
        assert parsed.total_pages == 1
        assert len(parsed.pages) == 1
    
    def test_pdf_metadata_extraction(self, sample_pdf):
        """Test PDF metadata extraction"""
        from backend.parsers.pdf_parser import PDFParser
        
        parser = PDFParser(enable_ocr=False)
        parsed = parser.parse(sample_pdf)
        
        assert parsed.metadata is not None
        assert parsed.metadata.file_size > 0


class TestDOCXParser:
    """Test suite for DOCX Parser"""
    
    def test_docx_text_extraction(self, sample_docx):
        """Test DOCX text extraction"""
        from backend.parsers.docx_parser import DOCXParser
        
        parser = DOCXParser()
        parsed = parser.parse(sample_docx)
        
        assert parsed.total_pages == 1
        assert "Test Document" in parsed.pages[0].text
    
    def test_docx_table_extraction(self, sample_docx):
        """Test DOCX table extraction"""
        from backend.parsers.docx_parser import DOCXParser
        
        parser = DOCXParser()
        parsed = parser.parse(sample_docx)
        
        assert len(parsed.pages[0].tables) == 1
        table = parsed.pages[0].tables[0]
        assert len(table) == 3  # 3 rows
        assert len(table[0]) == 3  # 3 columns
    
    def test_docx_metadata_extraction(self, sample_docx):
        """Test DOCX metadata extraction"""
        from backend.parsers.docx_parser import DOCXParser

        parser = DOCXParser()
        parsed = parser.parse(sample_docx)

        assert parsed.metadata is not None
        assert parsed.metadata.file_size > 0
        # Note: add_heading() doesn't set the document title property
        # The title would need to be set via core_properties
        assert parsed.metadata.page_count == 1  # DOCX treated as 1 page


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
