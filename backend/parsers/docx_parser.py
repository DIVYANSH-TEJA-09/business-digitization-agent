"""
DOCX Parser

Extracts text, tables, and images from DOCX files using python-docx.
"""
import os
import zipfile
import io
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image

from backend.models.schemas import ParsedDocument, Page, DocumentMetadata
from backend.models.enums import FileType
from backend.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):
    """
    DOCX parser with structure preservation
    
    Extracts:
    - Paragraphs with formatting
    - Tables
    - Embedded images
    """
    
    def __init__(self):
        """Initialize DOCX parser"""
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ParsedDocument object
        """
        self.validate_file(file_path)
        
        try:
            return self._parse_docx(file_path)
        except zipfile.BadZipFile:
            # Try old DOC format (limited support)
            return self._parse_doc_fallback(file_path)
    
    def _parse_docx(self, docx_path: str) -> ParsedDocument:
        """
        Parse DOCX file
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            ParsedDocument object
        """
        doc = Document(docx_path)
        
        # Extract elements
        elements = []
        full_text_parts = []
        tables = []
        
        # Iterate through document content
        for para in doc.paragraphs:
            if para.text.strip():
                para_data = self._parse_paragraph(para)
                elements.append({
                    'type': 'paragraph',
                    'data': para_data
                })
                full_text_parts.append(para.text)
        
        # Extract tables separately
        for table in doc.tables:
            table_data = self._parse_table(table)
            elements.append({
                'type': 'table',
                'data': table_data
            })
            tables.append(table_data)
            # Add table text to full text
            for row in table_data:
                full_text_parts.append(' | '.join(row))
        
        # Extract images
        images = self._extract_images(doc, docx_path)
        
        # Build full text
        full_text = '\n\n'.join(full_text_parts)
        
        # DOCX doesn't have pages, treat as single page
        page = Page(
            number=1,
            text=full_text,
            tables=tables,
            images=images,
            metadata={'elements_count': len(elements), 'paragraphs': len(doc.paragraphs), 'tables': len(doc.tables)}
        )
        
        return ParsedDocument(
            doc_id=self.generate_doc_id(docx_path),
            source_file=docx_path,
            file_type=FileType.DOCX,
            pages=[page],
            total_pages=1,
            metadata=self._extract_metadata(doc, docx_path),
            parsing_errors=[]
        )
    
    def _iter_block_items(self, doc: DocumentType):
        """
        Iterate through document block items (paragraphs and tables)
        
        Args:
            doc: python-docx Document object
            
        Yields:
            Paragraph or Table objects
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        
        for child in doc.element.body.iterchildren():
            if child.tag.endswith('tbl'):  # Table - check tables first
                yield Table(child, doc)
            elif child.tag.endswith('p'):  # Paragraph
                # Skip paragraphs inside tables (they're handled separately)
                parent_tbl = child.getparent()
                if parent_tbl is not None and parent_tbl.tag.endswith('tbl'):
                    continue
                yield Paragraph(child, doc)
    
    def _parse_paragraph(self, paragraph) -> Dict[str, Any]:
        """
        Parse paragraph with formatting
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            Paragraph data dictionary
        """
        return {
            'text': paragraph.text,
            'style': paragraph.style.name if paragraph.style else 'Normal',
            'is_heading': self._is_heading(paragraph),
            'alignment': str(paragraph.alignment) if paragraph.alignment else None,
            'runs_count': len(paragraph.runs)
        }
    
    def _is_heading(self, paragraph) -> bool:
        """
        Check if paragraph is a heading
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            True if heading
        """
        try:
            if not paragraph.style:
                return False
            
            style_name = paragraph.style.name.lower() if paragraph.style.name else ''
            return 'heading' in style_name or 'title' in style_name
        except Exception:
            return False
    
    def _parse_table(self, table) -> List[List[str]]:
        """
        Parse table to 2D array
        
        Args:
            table: python-docx Table object
            
        Returns:
            2D list of cell values
        """
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        
        return table_data
    
    def _extract_images(self, doc: DocumentType, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract embedded images from DOCX
        
        Args:
            doc: python-docx Document object
            docx_path: Path to DOCX file
            
        Returns:
            List of image metadata
        """
        images = []
        
        try:
            # DOCX files are ZIP archives
            with zipfile.ZipFile(docx_path) as docx_zip:
                # Images are in word/media/ folder
                media_files = [
                    f for f in docx_zip.namelist()
                    if f.startswith('word/media/')
                ]
                
                for i, media_file in enumerate(media_files):
                    try:
                        # Extract image bytes
                        image_bytes = docx_zip.read(media_file)
                        
                        # Determine format
                        image_format = self._detect_image_format(media_file)
                        
                        # Get dimensions
                        try:
                            with Image.open(io.BytesIO(image_bytes)) as img:
                                width, height = img.size
                        except Exception:
                            width, height = 0, 0
                        
                        images.append({
                            'image_id': f"docx_img_{i}",
                            'file_name': os.path.basename(media_file),
                            'width': width,
                            'height': height,
                            'file_size': len(image_bytes),
                            'mime_type': f"image/{image_format}",
                            'data': image_bytes  # Bytes for later use
                        })
                        
                    except Exception:
                        pass
                        
        except Exception:
            pass
        
        return images
    
    def _detect_image_format(self, filename: str) -> str:
        """
        Detect image format from filename
        
        Args:
            filename: Image filename
            
        Returns:
            Format string (jpeg, png, gif, etc.)
        """
        ext = os.path.splitext(filename)[1].lower()
        
        format_map = {
            '.jpg': 'jpeg',
            '.jpeg': 'jpeg',
            '.png': 'png',
            '.gif': 'gif',
            '.bmp': 'bmp',
            '.webp': 'webp',
        }
        
        return format_map.get(ext, 'unknown')
    
    def _extract_metadata(self, doc: DocumentType, docx_path: str) -> DocumentMetadata:
        """
        Extract DOCX metadata
        
        Args:
            doc: python-docx Document object
            docx_path: Path to file
            
        Returns:
            DocumentMetadata object
        """
        core_props = doc.core_properties
        
        return DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            creation_date=core_props.created,
            modification_date=core_props.modified,
            page_count=1,  # DOCX doesn't have fixed pages
            file_size=os.path.getsize(docx_path)
        )
    
    def _parse_doc_fallback(self, doc_path: str) -> ParsedDocument:
        """
        Fallback for old DOC format (limited support)
        
        Args:
            doc_path: Path to DOC file
            
        Returns:
            ParsedDocument object
        """
        # Try to read as plain text (best effort)
        try:
            with open(doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            return ParsedDocument(
                doc_id=self.generate_doc_id(doc_path),
                source_file=doc_path,
                file_type=FileType.DOC,
                pages=[Page(number=1, text=text, tables=[], images=[])],
                total_pages=1,
                metadata=DocumentMetadata(
                    file_size=os.path.getsize(doc_path)
                ),
                parsing_errors=['DOC format - limited parsing, text only']
            )
        except Exception as e:
            raise Exception(f"Failed to parse DOC file: {str(e)}")
